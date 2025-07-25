#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from num2words import num2words
import datetime
import subprocess
import os
from docx.shared import Pt, Inches
from find_court_uk import get_address_info
from datetime import datetime
import locale
import openpyxl
from openpyxl.styles import Font, Alignment
import json
from datetime import datetime

locale.setlocale(locale.LC_TIME, 'ru_RU.UTF-8')

ADDRESS_ABBREVIATIONS = {
    'проспект': 'пр-т',
    'Проспект': 'пр-т',
    'Переулок': 'пер.',
    'переулок': 'пер.',
    'Улица': 'ул.',
    'улица': 'ул.',
    'дом': 'д.',
    'Дом': 'д.',
    'квартира': 'кв.',
    'Квартира': 'кв.',
    'строение': 'стр.',
    'Строение': 'стр.',
    'Корпус': 'корп.',
    'корпус': 'корп.',
}

def load_data_from_json(json_data):
    """Преобразует данные из JSON в нужный формат для генерации документов"""
    # Преобразование периода
    start_date = datetime.strptime(
        f"01.{json_data['period']['from']['month']}.{json_data['period']['from']['year']}", 
        "%d.%m.%Y"
    )
    end_date = datetime.strptime(
        f"01.{json_data['period']['to']['month']}.{json_data['period']['to']['year']}", 
        "%d.%m.%Y"
    )
    
    # Формирование адреса
    street_name = json_data['address']['street']
    street_type = "пр-т" if is_prospekt(street_name) else "ул."
    
    full_address = f"{street_name}, {json_data['address']['house']}, {json_data['address']['aparts']}"
    address_parts = [
        f"г. Владивосток",
        f"{street_type} {street_name}",
        f"д. {json_data['address']['house']}",
        f"кв. {json_data['address']['aparts']}"
    ]
    
    # Получение информации о суде и УК по адресу
    address_data = get_address_info(
        street=json_data['address']['street'],
        house=json_data['address']['house'],
        appart=json_data['address']['aparts']
    )
    
    # Если найдены УК, используем первую из списка
    if address_data['uks']:
        uk_data = address_data['uks'][0]
    else:
        uk_data = {
            "uk": "Неизвестная УК",
            "number": "неизвестно",
            "location": "неизвестно",
            "inn": "неизвестно",
            "kpp": "неизвестно",
            "ogrn": "неизвестно",
            "reg_date": "неизвестно"
        }
    
    # Если найдены суды, используем первый из списка
    if address_data['courts']:
        court_data = address_data['courts'][0]
        court_number = court_data["court"].replace("Мировой суд №", "")
        court_district = court_data["region"]
    else:
        court_number = "17"
        court_district = "Первореченского"
    
    # Форматирование адреса для документа
    debtor_address_document = format_address_for_document(address_parts)
    
    # Форматирование адреса для имени файла
    address_parts_for_filename = [
        json_data['address']['street'],
        json_data['address']['house'],
        json_data['address']['aparts']
    ]
    address_for_filename = format_address_for_filename(address_parts_for_filename)
    
    # Подготовка данных для документа
    result_data = {
        "court_number": court_number,
        "court_district": court_district,
        "uk_name": uk_data['uk'],
        "uk_address": uk_data['location'],
        "debtor_address": full_address,
        "debtor_address_document": debtor_address_document,
        "address_for_filename": address_for_filename,
        "account_number": json_data['ca_number'],
        "debt_amount": json_data['total'],
        "penalty_amount": json_data['fine_total'],
        "state_duty": calculate_fee(json_data['total'], json_data['fine_total']),
        "period_str": format_date_range(start_date, end_date),
        "start_date": start_date,
        "end_date": end_date,
        "uk_inn": uk_data['inn'],
        "uk_kpp": uk_data['kpp'],
        "uk_ogrn": uk_data['ogrn'],
        "uk_reg_date": uk_data['reg_date'],
        "phone": "+7 9510008227",
        "representative": "Довгалюк Антон Владимирович",
        "raw_street": json_data['address']['street'],
        "raw_house": json_data['address']['house'],
        "appart": json_data['address']['aparts'],
        "street_type": street_type,  # Добавляем тип улицы
        "street_name": street_name   # Добавляем название улицы без типа
    }
    
    return result_data

def process_json_data(json_data):
    """Обрабатывает JSON-данные и создает документы"""
    try:
        # Преобразуем данные в нужный формат
        data = load_data_from_json(json_data)
        
        # Создаём документ Word
        doc = create_word_document(data)
        
        # Сохранение документов
        today = datetime.now().strftime("%Y%m%d_%H%M%S")
        os.makedirs("documents", exist_ok=True)
        docx_path = os.path.join("documents", f'Заявление_{data["address_for_filename"]}.docx')
        pdf_path = os.path.join("documents", f'Заявление_{data["address_for_filename"]}.pdf')
        doc.save(docx_path)

        # Конвертация в PDF
        try:
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                           os.path.dirname(pdf_path), docx_path], check=True)
            print(f"\nДокументы успешно созданы:\n{docx_path}\n{pdf_path}")
        except subprocess.CalledProcessError as e:
            print(f"\nОшибка конвертации. Убедитесь, что LibreOffice установлен и работает.\n"
                f"DOCX файл сохранён: {docx_path}\n"
                f"Ошибка: {e}")
        except Exception as e:
            print(f"\nНеизвестная ошибка: {e}\nDOCX файл сохранён: {docx_path}")
        
        return data
    
    except Exception as e:
        print(f"Ошибка обработки данных: {str(e)}")
        return None
    
def is_prospekt(street_name):
    """Проверяет, начинается ли название улицы с 'Проспект'"""
    return street_name.strip().startswith(('Проспект', 'проспект'))

def format_address_for_document(address_parts):
    formatted = []
    for part in address_parts:
        # Обрабатываем случай с проспектом
        if is_prospekt(part):
            part = part.replace("Проспект", "пр-т").replace("проспект", "пр-т")
        
        for full, short in ADDRESS_ABBREVIATIONS.items():
            if part.startswith(full):
                part = part.replace(full, short)
                break
        formatted.append(part)
    return ', '.join(formatted)

def format_address_for_filename(address_parts):
    parts = []
    for part in address_parts:
        clean_part = part
        for abbr in ADDRESS_ABBREVIATIONS.values():
            if part.startswith(abbr):
                clean_part = part.replace(abbr, '').strip()
                break
        parts.append(clean_part.replace(' ', '_').replace('/', '_'))
    return '_'.join(parts)

def format_date_range(start_date, end_date):
    month_cases = {
        'января': 'январь',
        'февраля': 'февраль',
        'марта': 'март',
        'апреля': 'апрель',
        'мая': 'май',
        'июня': 'июнь',
        'июля': 'июль',
        'августа': 'август',
        'сентября': 'сентябрь',
        'октября': 'октябрь',
        'ноября': 'ноябрь',
        'декабря': 'декабрь'
    }
    
    start_str = start_date.strftime("%B %Y года").lower()
    end_month_genitive = end_date.strftime("%B").lower()
    end_month_nominative = month_cases.get(end_month_genitive, end_month_genitive)
    end_str = f"по {end_month_nominative} {end_date.year} года"
    
    return f"{start_str} {end_str}"

def parse_address(full_address):
    """Парсер адреса с использованием нормализованного формата"""
    normalized = normalize_address(full_address)
    parts = [p.strip() for p in normalized.split(',')]
    
    street = parts[0] if len(parts) > 0 else ''
    house_info = parts[1] if len(parts) > 1 else ''
    apartment = parts[2] if len(parts) > 2 else ''
    
    # Определяем тип улицы
    street_type = "пр-т" if is_prospekt(street) else "ул."
    street_name = street.replace("Проспект", "").replace("проспект", "").strip() if is_prospekt(street) else street
    
    # Удаляем "дом" и "квартира" из соответствующих полей
    house_info = house_info.replace('дом', '').strip()
    apartment = apartment.replace('квартира', '').strip()
    
    return {
        'street': street_name,
        'street_type': street_type,
        'house_info': f"дом {house_info}" if house_info else '',
        'apartment': f"кв. {apartment}" if apartment else ''
    }

def normalize_address(address_str):
    """Приводит адрес к единому формату для поиска в базе данных"""
    normalized = ' '.join(address_str.lower().split())
    
    replacements = {
        'пр-т': 'проспект',
        'пр.': 'проспект',
        'пер.': 'переулок',
        'ул.': 'улица',
        'д.': 'дом',
        'кв.': 'квартира',
        'стр.': 'строение',
        'корп.': 'корпус',
        'г.': 'город',
        'г ': 'город ',
        'владивосток': 'город владивосток'
    }
    
    for short, long in replacements.items():
        normalized = normalized.replace(short, long)
    
    normalized = normalized.replace('город владивосток', 'владивосток')
    normalized = ' '.join(word.capitalize() for word in normalized.split())
    
    return normalized

def format_money_short(amount):
    rub = int(amount)
    kop = int(round((amount - rub) * 100))
    
    num_format = f"{rub:,d}".replace(",", " ")
    
    last_digit = rub % 10
    last_two = rub % 100
    
    if last_two in range(11, 20):
        rub_end = 'рублей'
    else:
        if last_digit == 1:
            rub_end = 'рубль'
        elif last_digit in [2,3,4]:
            rub_end = 'рубля'
        else:
            rub_end = 'рублей'
    
    if kop > 0:
        last_digit_kop = kop % 10
        last_two_kop = kop % 100
        
        if last_two_kop in range(11, 20):
            kop_end = 'копеек'
        else:
            if last_digit_kop == 1:
                kop_end = 'копейка'
            elif last_digit_kop in [2,3,4]:
                kop_end = 'копейки'
            else:
                kop_end = 'копеек'
    else:
        kop_end = 'копеек'
    
    return f"{num_format} {rub_end} {kop:02d} {kop_end}"

def format_money_long(amount):
    rub = int(amount)
    kop = int(round((amount - rub) * 100))
    
    num_format = f"{rub:,d}".replace(",", " ") + f",{kop:02d}"
    
    rub_word = num2words(rub, lang='ru')
    kop_word = num2words(kop, lang='ru') if kop > 0 else ''
    
    last_digit = rub % 10
    last_two = rub % 100
    
    if last_two in range(11, 20):
        rub_end = 'рублей'
    else:
        if last_digit == 1:
            rub_end = 'рубль'
        elif last_digit in [2,3,4]:
            rub_end = 'рубля'
        else:
            rub_end = 'рублей'
    
    words = f"{rub_word} {rub_end}"
    if kop > 0:
        words += f" {kop:02d} коп."
    else:
        words += " 00 коп."
    
    return f"{num_format} ({words})"

def calculate_fee(debt, penalty):
    price = debt + penalty
    if price <= 100_000:
        return 4_000 / 2
    elif 100_000 < price <= 300_000:
        return 4_000 + 0.03 * (price - 100_000) / 2
    elif 300_000 < price <= 500_000:
        return 10_000 + 0.025 * (price - 300_000) / 2
    else:
        return 0

def set_red_line(paragraph, indent_size=0.5):
    paragraph.paragraph_format.first_line_indent = Inches(indent_size / 2.54)

def add_simple_signature(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("_________________________")
    p.add_run("Довгалюк Антон Владимирович")

def create_excel_report(data_list, filename):
    file_exists = os.path.exists(os.path.join("reports", filename))
    
    if file_exists:
        wb = openpyxl.load_workbook(os.path.join("reports", filename))
        ws = wb.active
        
        if ws.max_row == 1:
            headers = [
                "Улица", "Дом", "Квартира", "Лицевой счёт", 
                "Начало периода", "Конец периода", "Сумма основного долга", 
                "Сумма пени", "Сумма пошлины", "Управляющая компания",
                "Номер судебного участка", "Район суда", "Статус"
            ]
            ws.append(headers)
            
            for col in range(1, len(headers) + 1):
                cell = ws.cell(row=1, column=col)
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center')
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Судебные приказы"
        
        headers = [
            "Улица", "Дом", "Квартира", "Лицевой счёт", 
            "Начало периода", "Конец периода", "Сумма основного долга", 
            "Сумма пени", "Сумма пошлины", "Управляющая компания",
            "Номер судебного участка", "Район суда"
        ]
        ws.append(headers)
        
        for col in range(1, len(headers) + 1):
            cell = ws.cell(row=1, column=col)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
    
    for data in data_list:
        addr_parts = data['debtor_address'].split(', ')
        street = addr_parts[0] if len(addr_parts) > 0 else ''
        house = addr_parts[1] if len(addr_parts) > 1 else ''
        apartment = addr_parts[2] if len(addr_parts) > 2 else ''
        
        uk_name = data.get('uk_name', 'Неизвестно')
        court_number = data.get('court_number', 'Неизвестно')
        court_district = data.get('court_district', 'Неизвестно')
        
        row = [
            street,
            house,
            apartment,
            data['account_number'],
            data['start_date'].strftime("%d.%m.%Y"),
            data['end_date'].strftime("%d.%m.%Y"),
            data['debt_amount'],
            data['penalty_amount'],
            data['state_duty'],
            uk_name,
            court_number,
            court_district,
        ]
        ws.append(row)
    
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2) * 1.2
        ws.column_dimensions[column_letter].width = adjusted_width
    
    os.makedirs("reports", exist_ok=True)
    report_path = os.path.join("reports", filename)
    wb.save(report_path)
    print(f"\nExcel-отчёт сохранён: {report_path}")


def create_word_document(data):
    doc = Document()
    
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.6)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(8.95)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    total = data["debt_amount"] + data["penalty_amount"] + data["state_duty"]
    addr = parse_address(data["debtor_address"])

    debt_short = format_money_short(data["debt_amount"])
    penalty_short = format_money_short(data["penalty_amount"])
    duty_short = f"{int(data['state_duty']):,d}".replace(",", " ") + " рублей"
    total_short = format_money_short(total)

    debt_long = format_money_long(data["debt_amount"])
    penalty_long = format_money_long(data["penalty_amount"])
    duty_long = format_money_long(data["state_duty"])

    def add_right_block(text, bold=False, italic=False, indent=Inches(3.2)):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = indent
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p

    add_right_block(f'Мировому судье Судебного участка № {data["court_number"]}', bold=True)
    add_right_block(f'{data["court_district"]} судебного района', bold=True)
    add_right_block('г. Владивосток', bold=True)
    doc.add_paragraph()

    add_right_block('Заявитель:', bold=True, italic=True)
    add_right_block(f'{data["uk_name"]}')
    add_right_block('Местонахождение Заявителя:', bold=True, italic=True)
    add_right_block(f"{data['uk_address']}")
    doc.add_paragraph()

    add_right_block(f'телефон: {data["phone"]}', bold=True)
    add_right_block('Должник:', bold=True, italic=True)
    add_right_block('Должник не установлен', bold=True)
    doc.add_paragraph()

    add_right_block('Год рождения: данные неизвестны')
    add_right_block('Место рождения: Неизвестно, идентификационные')
    add_right_block('данные отсутствуют')

    doc.add_paragraph()
    add_right_block('Адрес места жительства:', bold=True)
    add_right_block(f'{data["debtor_address_document"].replace("Проспект", "")}')

    def add_money_line(label, amount_str):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(3.2)
        run = p.add_run(f"{label}: {amount_str}")
        run.bold = True
        run.italic = True

    doc.add_paragraph()
    add_money_line('Сумма основного долга', debt_short)
    doc.add_paragraph()
    add_money_line('Сумма пени', penalty_short)
    doc.add_paragraph()
    add_money_line('Государственная пошлина', duty_short)
    doc.add_paragraph()
    add_money_line('Всего к взысканию', total_short)
    doc.add_paragraph()

    p = doc.add_paragraph('ЗАЯВЛЕНИЕ')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p = doc.add_paragraph('О выдаче судебного приказа')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True

    main_text = [
        f'{data["uk_name"]} занимается управлением, содержанием и ремонтом общего имущества дома № {addr["house_info"].replace("дом", "").strip()} по {data["street_type"]} {addr["street"]} в г. Владивостоке, что подтверждается протоколом о выборе способа управления многоквартирным домом, договором управления многоквартирным домом, утвержденным решением общего собрания и обязательным для всех собственников помещений.',
        
        f'Плательщик за содержание и ремонт жилья квартиры № {addr["apartment"].replace("кв.", "").strip()} в доме №{addr["house_info"].replace("дом", "").strip()} по {data["street_type"]} {addr["street"]} не установлен. У заявителя информация о собственнике помещения отсутствует, а также отсутствует возможность самостоятельно запросить в регистрирующих органах данные сведения.',
        
        f'У Ответчика имеется задолженность перед {data["uk_name"]} в размере {debt_long}',
        
        'В соответствии со ст. 153, 155 ЖК РФ граждане обязаны своевременно и полностью вносить плату за жилое помещение и коммунальные услуги ежемесячно до 10 числа месяца, следующего за истекшим.',
        
        f'Образовавшаяся за период c {data["period_str"]} задолженность составляет {debt_long}, а также сумма пени за вышеуказанный период составляет {penalty_long}. На основании вышеизложенного и руководствуясь ст. 121- 124, 131 ГПК РФ'
    ]

    for text in main_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if "ст." in text and "РФ" in text and "121" in text:
            start = text.find("ст.")
            end = text.find("РФ") + len("РФ")
            
            p.add_run(text[:start])
            p.add_run(text[start:end]).bold = True
            p.add_run(text[end:])
        else:
            p.add_run(text)
        
        set_red_line(p)
        p.paragraph_format.space_after = Pt(6)
        
    p = doc.add_paragraph('ПРОШУ:')
    p.runs[0].bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    request_text = [
        f'Запросить сведения о собственнике помещения и его регистрации. Выдать судебный приказ о взыскании с Должника, данные неизвестны, в пользу {data["uk_name"]}, ИНН {data["uk_inn"]}, КПП {data["uk_kpp"]}, ОГРН {data["uk_ogrn"]}, дата гос. регистрации {data["uk_reg_date"]} задолженности по оплате за содержание и ремонт жилья в размере: {debt_long}, сумму пени в размере: {penalty_long}.',
        
        f'Включить в судебный приказ о взыскании с Должника, данные неизвестны, в пользу {data["uk_name"]}, ИНН {data["uk_inn"]}, КПП {data["uk_kpp"]}, ОГРН {data["uk_ogrn"]}, дата гос. регистрации {data["uk_reg_date"]} сумму государственной пошлины в размере: {duty_long}.',
    ]

    for text in request_text:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_red_line(p)
        p.paragraph_format.space_after = Pt(6)

    state = [
        'В соответствии со ст. 333.40 НК РФ выдать справку о возврате государственной пошлины, если адрес регистрации Должника не входит в состав Владивостокского Городского округа.'
    ]
    for text in state:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.runs[0].bold = True

    p = doc.add_paragraph('Приложение:')
    p.runs[0].bold = True
    
    attachments = [
        'Расчет суммы задолженности',
        'Реестр начисления пени',
        'Копия протокола о выборе способа управления домом',
        'Копия доверенности представителя',
        'Копия квитанции об оплате пошлины'
    ]

    for item in attachments:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()
    p = doc.add_paragraph('Представитель (по доверенности)')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].italic = True
    
    p = doc.add_paragraph(f'{data["uk_name"]}')
    doc.add_paragraph()
    doc.add_paragraph()
    add_simple_signature(doc)
    
    return doc

def generate_documents(json_data, output_dir="documents"):
    """Генерирует все документы (Word, PDF, Excel) для одного дела"""
    try:
        # Создаем папку для документов, если ее нет
        os.makedirs(output_dir, exist_ok=True)
        
        # Преобразуем данные в нужный формат
        data = load_data_from_json(json_data)
        
        # Генерируем Word документ
        doc = create_word_document(data)
        
        # Сохраняем документы
        today = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f'Заявление_{data["address_for_filename"]}'
        
        # Сохраняем Word
        docx_path = os.path.join(output_dir, f'{base_filename}.docx')
        doc.save(docx_path)
        
        # Конвертируем в PDF
        pdf_path = os.path.join(output_dir, f'{base_filename}.pdf')
        try:
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                           output_dir, docx_path], check=True)
        except subprocess.CalledProcessError:
            print(f"Ошибка конвертации в PDF. Убедитесь, что LibreOffice установлен.")
        
        return {
            'docx_path': docx_path,
            'pdf_path': pdf_path,
            'data': data
        }
    
    except Exception as e:
        print(f"Ошибка генерации документов: {e}")
        return None

def generate_excel_report(data_list, output_dir="reports"):
    """Генерирует Excel отчет по списку дел"""
    try:
        os.makedirs(output_dir, exist_ok=True)
        today = datetime.now().strftime("%Y%m%d")
        excel_filename = f"Судебные_приказы_{today}.xlsx"
        excel_path = os.path.join(output_dir, excel_filename)
        
        create_excel_report(data_list, excel_filename)
        
        return excel_path
    except Exception as e:
        print(f"Ошибка генерации Excel отчета: {str(e)}")
        return None
    
if __name__ == "__main__":    
    cases = []
    
    # Получаем данные из внешнего источника
    debts_data = get_debts_data()
    
    for json_data in debts_data:
        # Обрабатываем каждое дело из списка
        try:
            # Преобразуем данные в нужный формат
            data = load_data_from_json(json_data)
            
            # Создаём документ Word
            doc = create_word_document(data)
            
            # Сохранение документов
            today = datetime.now().strftime("%Y%m%d_%H%M%S")
            os.makedirs("documents", exist_ok=True)
            docx_path = os.path.join("documents", f'Заявление_{data["address_for_filename"]}.docx')
            pdf_path = os.path.join("documents", f'Заявление_{data["address_for_filename"]}.pdf')
            doc.save(docx_path)

            # Конвертация в PDF
            try:
                subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                               os.path.dirname(pdf_path), docx_path], check=True)
                print(f"\nДокументы успешно созданы:\n{docx_path}\n{pdf_path}")
            except subprocess.CalledProcessError as e:
                print(f"\nОшибка конвертации. Убедитесь, что LibreOffice установлен и работает.\n"
                    f"DOCX файл сохранён: {docx_path}\n"
                    f"Ошибка: {e}")
            except Exception as e:
                print(f"\nНеизвестная ошибка: {e}\nDOCX файл сохранён: {docx_path}")
            
            # Добавляем данные в список для Excel-отчёта
            cases.append(data)
            
        except Exception as e:
            print(f"Ошибка обработки данных: {str(e)}")
            continue
    
    # Создаём сводный Excel-отчёт
    if cases:
        today = datetime.now().strftime("%Y%m%d")
        excel_filename = f"Судебные_приказы_{today}.xlsx"
        create_excel_report(cases, excel_filename)
